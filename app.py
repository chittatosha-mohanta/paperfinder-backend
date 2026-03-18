from groq import Groq
import os
import json
import base64
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re
from flask import Flask, request, jsonify
from flask_cors import CORS
import fitz
import pdfplumber

app = Flask(__name__)
CORS(app, origins=[
    "http://localhost:5173",
    "http://localhost:5174",
    "https://paperfinder-pro.vercel.app"
])

# ── Helper functions ─────────────────────────────────────────

def clean_text(text):
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def sanitize_text(text):
    """
    Replace common Unicode characters that break python-docx XML serialization
    with safe ASCII equivalents. Also strips NULL bytes and control characters.
    """
    if not text:
        return text
    replacements = {
        '\u2018': "'",   # left single quote
        '\u2019': "'",   # right single quote / smart apostrophe
        '\u201c': '"',   # left double quote
        '\u201d': '"',   # right double quote
        '\u2013': '-',   # en dash
        '\u2014': '--',  # em dash (will be re-added as real em dash where needed)
        '\u2026': '...',  # ellipsis
        '\u00a0': ' ',   # non-breaking space
        '\u00ad': '-',   # soft hyphen
        '\u200b': '',    # zero-width space
        '\u200c': '',    # zero-width non-joiner
        '\u200d': '',    # zero-width joiner
        '\ufeff': '',    # BOM
        '\u2022': '-',   # bullet
        '\u2012': '-',   # figure dash
        '\u2015': '-',   # horizontal bar
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Strip NULL bytes and control characters (keep newline \n, tab \t)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text

def extract_query(text):
    stop_words = {
        'the','and','for','with','from','this','that','are','was','were','has',
        'have','been','will','can','may','our','their','also','which','when',
        'where','how','what','all','not','but','its','than','more','some','such',
        'each','both','only','very','most','into','over','after','under','about',
        'other','these','those','then','them','they','would','could','should',
        'while','between','through','during','before','any','use','used','using',
        'based','show','shows','paper','study','research','article','results',
        'result','method','methods','data','model','models','figure','table',
        'section','university','college','china','shanghai','received','accepted',
        'published','copyright','journal','volume','correspondence','academic',
        'address','revised','april','march','october','january','february','june',
        'july','august','september','november','december','email','doi','http',
        'www','page','pages'
    }
    words = text.split()
    keywords = [w for w in words if len(w) > 4 and w.isalpha() and w.lower() not in stop_words]
    seen = set()
    unique_keywords = []
    for w in keywords:
        lower = w.lower()
        if lower not in seen:
            seen.add(lower)
            unique_keywords.append(w)
        if len(unique_keywords) == 12:
            break
    return ' '.join(unique_keywords)

def extract_with_pymupdf(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for page_num in range(min(2, len(doc))):
        text += doc[page_num].get_text() + " "
    doc.close()
    return clean_text(text)

def extract_with_pdfplumber(pdf_bytes):
    text = ""
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages[:2]:
            page_text = page.extract_text()
            if page_text:
                text += page_text + " "
    return clean_text(text)

# ── Document helpers ──────────────────────────────────────────

def set_run_font(run, size_pt, bold=False, italic=False, font_name='Times New Roman'):
    """Apply Times New Roman font to a run."""
    run.font.name  = font_name
    run.font.size  = Pt(size_pt)
    run.bold       = bold
    run.italic     = italic

def set_two_columns(section):
    """Apply two-column layout to a section."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:cols')):
        sectPr.remove(existing)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'),        '2')
    cols.set(qn('w:space'),      '720')
    cols.set(qn('w:equalWidth'), '1')
    sectPr.append(cols)

def set_single_column(section):
    """Explicitly set single-column layout on a section."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:cols')):
        sectPr.remove(existing)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '1')
    sectPr.append(cols)

def insert_section_break_single_to_double(doc):
    """
    Insert a continuous section break marking the END of the single-column
    header area. The sectPr inside this paragraph describes Section 0
    (header = 1 col, with explicit pgSz + pgMar so the first page renders
    correctly in all Word versions and PDF export).

    Structure after this:
      Section 0 (continuous, 1-col) = title, abstract, index terms
      Section 1 (body, 2-col)       = all body paragraphs
    """
    para = doc.add_paragraph()
    pf   = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing = Pt(1)

    pPr    = para._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')

    # Continuous break
    pgType = OxmlElement('w:type')
    pgType.set(qn('w:val'), 'continuous')
    sectPr.append(pgType)

    # Single column for the header area
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '1')
    sectPr.append(cols)

    # FIX: Explicitly define page size (US Letter) for Section 0
    # Without this, the first page has no defined dimensions and
    # renders incorrectly in some Word versions and PDF exports
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '12240')   # 8.5 inches in twips
    pgSz.set(qn('w:h'), '15840')   # 11 inches in twips
    sectPr.append(pgSz)

    # FIX: Explicitly define margins for Section 0
    # top=1440 (1in), bottom=1440 (1in), left=1083 (~0.75in), right=1083 (~0.75in)
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'),    '1440')
    pgMar.set(qn('w:right'),  '1083')
    pgMar.set(qn('w:bottom'), '1440')
    pgMar.set(qn('w:left'),   '1083')
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)

    pPr.append(sectPr)

def clean_line(line):
    """Strip markdown formatting from a line."""
    line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
    line = re.sub(r'\*(.+?)\*',     r'\1', line)
    line = re.sub(r'#{1,6}\s*',     '',    line)
    return line.strip()

def is_ieee_main_heading(line):
    """Detect IEEE main headings: I. INTRODUCTION, II. RELATED WORK, REFERENCES, etc."""
    s = line.strip().upper()
    if re.match(r'^[IVX]+\.\s+[A-Z][A-Z\s]+$', s):
        return True
    if s in ('REFERENCES', 'ACKNOWLEDGMENT', 'ACKNOWLEDGMENTS',
             'DATA AVAILABILITY', 'CONFLICTS OF INTEREST'):
        return True
    return False

def is_ieee_subsection_heading(line):
    """Detect IEEE subsection headings: A. Name, B. Another Name"""
    s = line.strip()
    return bool(re.match(r'^[A-Z]\.\s+[A-Z][a-zA-Z]', s)) and len(s) < 80

def is_generic_heading(line, keywords):
    """Fallback heading detector for non-IEEE formats."""
    s     = line.strip()
    lower = s.lower()
    clean = re.sub(r'^[IVXivx]+\.\s*|^\d+[\.\s]+', '', lower).strip().rstrip('.')
    return any(clean.startswith(kw) for kw in keywords) and len(s) < 90

def is_reference_line(line):
    return bool(re.match(r'^\[\d+\]', line.strip()))

def is_figure_placeholder(line):
    upper = line.upper()
    return '[DIAGRAM_HERE' in upper or '[FIGURE' in upper

# ── Routes ───────────────────────────────────────────────────

@app.route('/health', methods=['GET'])
def health():
    return jsonify({"status": "ok", "message": "PaperFinder backend running"})


@app.route('/extract-pdf', methods=['POST'])
def extract_pdf():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files['file']
    if not file.filename.endswith('.pdf'):
        return jsonify({"error": "File must be a PDF"}), 400
    pdf_bytes = file.read()
    try:
        text = extract_with_pymupdf(pdf_bytes)
        if text and len(text.split()) > 15:
            return jsonify({"success": True, "query": extract_query(text),
                            "method": "pymupdf", "preview": text[:200]})
    except Exception as e:
        print(f"PyMuPDF failed: {e}")
    try:
        text = extract_with_pdfplumber(pdf_bytes)
        if text and len(text.split()) > 15:
            return jsonify({"success": True, "query": extract_query(text),
                            "method": "pdfplumber", "preview": text[:200]})
    except Exception as e:
        print(f"pdfplumber failed: {e}")
    return jsonify({"error": "Could not extract text from this PDF."}), 422


@app.route('/extract-pdf-full', methods=['POST'])
def extract_pdf_full():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files['file']
    if not file.filename.endswith('.pdf'):
        return jsonify({"error": "File must be a PDF"}), 400
    pdf_bytes = file.read()
    try:
        ref_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text    = ""
        for page in ref_doc:
            text += page.get_text() + "\n"
        pages = len(ref_doc)
        ref_doc.close()
        text = clean_text(text)
        if not text or len(text.split()) < 10:
            raise Exception("Could not extract text from this PDF")
        return jsonify({"success": True, "text": text[:15000], "pages": pages})
    except Exception as e:
        return jsonify({"error": str(e)}), 422


@app.route('/chat-pdf', methods=['POST'])
def chat_pdf():
    question = request.form.get('question', '').strip()
    pdf_text  = request.form.get('pdf_text',  '').strip()
    if not question:
        return jsonify({"error": "No question provided"}), 400
    if not pdf_text:
        return jsonify({"error": "No PDF text provided"}), 400

    prompt = f"""You are a helpful research assistant. A user has uploaded a PDF and wants to ask questions about it.

PDF CONTENT:
{pdf_text[:8000]}

USER QUESTION:
{question}

INSTRUCTIONS:
- Answer ONLY based on the PDF content above
- Be specific and reference relevant sections when helpful
- If the answer is not found in the PDF, clearly say so
- Keep answers clear, concise and academic
- Format your answer with proper paragraphs
- If the question asks for a summary, provide a structured summary

Answer:"""

    try:
        client = Groq(api_key=os.environ.get('GROQ_API_KEY'))
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}]
        )
        answer = response.choices[0].message.content
        return jsonify({"success": True, "answer": answer})
    except Exception as e:
        return jsonify({"error": f"AI failed: {str(e)}"}), 500


@app.route('/generate-paper', methods=['POST'])
def generate_paper():
    title        = sanitize_text(request.form.get('title', ''))
    abstract     = sanitize_text(request.form.get('abstract', ''))
    paper_format = request.form.get('format', 'IEEE')

    # ── Read reference PDFs (max 5) ───────────────────────────
    reference_texts = []
    for key in sorted(request.files.keys()):
        if key.startswith('reference_') and len(reference_texts) < 5:
            try:
                pdf_bytes = request.files[key].read()
                ref_doc   = fitz.open(stream=pdf_bytes, filetype="pdf")
                text      = "".join(page.get_text() for page in ref_doc)
                ref_doc.close()
                reference_texts.append(clean_text(text)[:3000])
            except Exception as e:
                print(f"Failed to read reference: {e}")

    # ── Read diagram images (max 10) ──────────────────────────
    # FIX: diagrams were being uploaded but completely ignored before
    diagram_images = []   # list of (filename, bytes, extension)
    for key in sorted(request.files.keys()):
        if key.startswith('diagram_') and len(diagram_images) < 10:
            try:
                file      = request.files[key]
                img_bytes = file.read()
                filename  = file.filename or key
                ext       = filename.rsplit('.', 1)[-1].lower() if '.' in filename else 'png'
                if ext not in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff'):
                    ext = 'png'
                diagram_images.append((filename, img_bytes, ext))
            except Exception as e:
                print(f"Failed to read diagram: {e}")

    refs_summary = "\n\n".join(
        f"Reference {i+1}:\n{t}" for i, t in enumerate(reference_texts)
    ) if reference_texts else "No references provided."

    diagram_note = ""
    if diagram_images:
        diagram_note = f"\n\nDIAGRAMS PROVIDED: The user has uploaded {len(diagram_images)} diagram image(s). Place [DIAGRAM_HERE: Fig. N - describe what this figure shows] at appropriate locations in the paper to indicate where each diagram should appear (Fig. 1, Fig. 2, etc.)."

    # ── Format rules ──────────────────────────────────────────
    format_rules = {

        "IEEE": """STRICT IEEE JOURNAL FORMAT — FOLLOW EXACTLY:

OUTPUT STRUCTURE (write in this exact order):

━━━ ABSTRACT ━━━
Start the very first line with exactly:
Abstract—
(the word Abstract followed IMMEDIATELY by an em dash —, then the text on the SAME LINE)
Example: Abstract—This paper proposes a novel framework for deep learning...
- Single paragraph, 150-250 words, NO citations, NO bullet points

━━━ INDEX TERMS ━━━
Immediately after abstract on its own line:
Index Terms—term one, term two, term three, term four.

━━━ SECTION HEADINGS ━━━
I. INTRODUCTION
II. RELATED WORK
III. METHODOLOGY
IV. RESULTS AND DISCUSSION
V. CONCLUSION
REFERENCES
(Roman numeral + period + space + ALL CAPS — no markdown)

━━━ SUBSECTION HEADINGS ━━━
A. Subsection Name
B. Another Subsection
(Letter + period + space + Title Case)

━━━ BODY TEXT ━━━
- Full academic paragraphs only — NEVER bullet points
- In-text citations: [1], [2], [1]-[3]
- Figure placeholders: [DIAGRAM_HERE: Fig. N description]
- Equations on own line: expression   (1)

━━━ REFERENCES ━━━
[1] A. B. Author, "Title," Journal, vol. X, no. Y, pp. ZZ-ZZ, Year.
[2] A. B. Author, Title of Book. City: Publisher, Year.
[3] A. B. Author, "Title," in Proc. Conf., City, Year, pp. ZZ-ZZ.

RULES: No markdown, no bullets, no "Abstract:", minimum 2500 words, third person.""",

        "APA": """APA 7th Edition: Abstract bold label, Keywords: word1, word2.
Level 1 headings bold centered title case. In-text: (Author, Year).
References alphabetical hanging indent.
Journal: Author, A. B. (Year). Title. Journal, vol(issue), pages.
No bullets, 2500 words minimum.""",

        "ACM": """ACM Format: ABSTRACT all caps. CCS CONCEPTS and KEYWORDS after abstract.
Sections: 1. INTRODUCTION  2. RELATED WORK  3. METHODOLOGY  4. RESULTS  5. CONCLUSION
Subsections: 3.1 Name. Citations: [1], [2].
Ref: [1] Author. Year. Title. In Proceedings. ACM, pages.
No bullets, 2500 words minimum.""",

        "MLA": """MLA 9th Edition: Header block Author/Course/Instructor/Date. Title centered.
Headings centered bold title case. In-text: (Author page).
Works Cited alphabetical: Author Last, First. "Title." Journal, vol., Year, pp.
No bullets, 2500 words minimum.""",

        "Nature": """Nature Format: No abstract label, start paragraph directly, 150 words max.
Sections not numbered: Introduction / Results / Discussion / Methods / References
Superscript citations: word¹ — NOT [1].
Ref: Author, A. B. et al. Title. Journal vol, pages (year).
No bullets, 2500 words minimum.""",

        "Springer": """Springer Format: Abstract bold label, Keywords: word1 · word2 · word3
Sections: 1 Introduction  2 Related Work (number space Title Case, no period)
Subsections: 2.1 Name. Citations: [1], [2].
Ref: 1. Author, A.B.: Title. Journal 45(3), 123-145 (2020).
No bullets, 2500 words minimum."""
    }

    rules = format_rules.get(paper_format, format_rules["IEEE"])

    prompt = f"""You are an expert academic paper writer. Write a complete, professional research paper in {paper_format} format.

Paper Title: {title}

Abstract provided by author:
{abstract}

Reference papers to cite:
{refs_summary}{diagram_note}

{rules}

Start immediately with the Abstract. Do NOT write the title again. Do NOT add preamble.
Do NOT use markdown symbols (no **, no *, no ##).
Minimum 2500 words. All body text in flowing academic paragraphs, no bullet points."""

    try:
        client = Groq(api_key=os.environ.get('GROQ_API_KEY'))
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=6000,
            messages=[{"role": "user", "content": prompt}]
        )
        paper_content = sanitize_text(response.choices[0].message.content)
    except Exception as e:
        return jsonify({"error": f"AI generation failed: {str(e)}"}), 500

    # ════════════════════════════════════════════════════════
    # BUILD .docx
    # ════════════════════════════════════════════════════════
    try:
        doc = Document()

        # US Letter, IEEE margins
        for sec in doc.sections:
            sec.top_margin    = Cm(2.54)
            sec.bottom_margin = Cm(2.54)
            sec.left_margin   = Cm(1.91)
            sec.right_margin  = Cm(1.91)
            sec.page_width    = Cm(21.59)
            sec.page_height   = Cm(27.94)

        # ── Parse AI output ───────────────────────────────────
        lines          = paper_content.split('\n')
        abstract_lines = []
        keyword_lines  = []
        body_lines     = []
        mode           = 'before'

        body_triggers = [
            'i. ', 'ii. ', 'iii. ', 'iv. ', 'v. ',
            'introduction', 'related', 'methodology',
            '1. ', '2. ', '1 ', '2 '
        ]
        body_headings = [
            'introduction', 'related work', 'literature review',
            'methodology', 'proposed', 'method', 'approach',
            'results', 'discussion', 'experiment', 'evaluation',
            'conclusion', 'future work', 'references',
            'acknowledgment', 'acknowledgments',
            'data availability', 'conflicts of interest'
        ]

        for raw_line in lines:
            s     = raw_line.strip()
            if not s:
                continue
            lower = s.lower()

            if re.match(r'^abstract[—\-:]*\s*', lower) and mode == 'before':
                mode  = 'abstract'
                after = re.sub(r'^abstract[—\-:]*\s*', '', s, flags=re.IGNORECASE).strip()
                if after:
                    abstract_lines.append(after)
                continue

            if re.match(r'^(index terms?|keywords?)[—\-:]*', lower):
                mode = 'keywords'
                keyword_lines.append(s)
                continue

            if mode == 'abstract':
                is_body = (
                    is_ieee_main_heading(s) or
                    is_ieee_main_heading(s.upper()) or
                    (any(lower.startswith(t) for t in body_triggers) and len(s) < 90)
                )
                if is_body:
                    mode = 'body'
                    body_lines.append(s)
                else:
                    abstract_lines.append(s)

            elif mode == 'keywords':
                is_body = (
                    is_ieee_main_heading(s) or
                    is_ieee_main_heading(s.upper()) or
                    (any(lower.startswith(t) for t in body_triggers) and len(s) < 90)
                )
                if is_body:
                    mode = 'body'
                    body_lines.append(s)
                else:
                    keyword_lines.append(s)

            elif mode == 'body':
                body_lines.append(s)

        # ════════════════════════════════════════════════════
        # WRITE DOCUMENT — single column header area first
        # ════════════════════════════════════════════════════

        # Title
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(6)
        tr = tp.add_run(title)
        set_run_font(tr, size_pt=16, bold=True)

        # Author / affiliation placeholder (IEEE style)
        auth_para = doc.add_paragraph()
        auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_para.paragraph_format.space_before = Pt(4)
        auth_para.paragraph_format.space_after  = Pt(10)
        set_run_font(auth_para.add_run("Author Name — Institution Name"), size_pt=10)

        # Abstract
        abs_text = ' '.join(abstract_lines).strip()
        abs_text = re.sub(r'^abstract[—\-:]*\s*', '', abs_text, flags=re.IGNORECASE).strip()
        abs_text = re.sub(r'\s+', ' ', abs_text)
        if not abs_text:
            abs_text = abstract

        if paper_format == "IEEE":
            # All on ONE paragraph: bold "Abstract—" + normal body text
            abs_para = doc.add_paragraph()
            abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            abs_para.paragraph_format.space_before = Pt(0)
            abs_para.paragraph_format.space_after  = Pt(4)
            r_label = abs_para.add_run("Abstract")
            set_run_font(r_label, size_pt=9, bold=True)
            r_dash  = abs_para.add_run("\u2014")
            set_run_font(r_dash,  size_pt=9, bold=True)
            r_body  = abs_para.add_run(abs_text)
            set_run_font(r_body,  size_pt=9, bold=False, italic=False)
        else:
            ah = doc.add_paragraph()
            ah.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ah.paragraph_format.space_after = Pt(2)
            set_run_font(ah.add_run("Abstract"), size_pt=10, bold=True)
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            ap.paragraph_format.space_after = Pt(6)
            set_run_font(ap.add_run(abs_text), size_pt=10, italic=True)

        # Index Terms / Keywords
        kw_raw = ' '.join(keyword_lines).strip()
        kw_raw = re.sub(r'\s+', ' ', kw_raw)
        if kw_raw:
            kp = doc.add_paragraph()
            kp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            kp.paragraph_format.space_before = Pt(0)
            kp.paragraph_format.space_after  = Pt(10)
            if paper_format == "IEEE":
                terms = re.sub(r'^(index terms?|keywords?)[—\-:]*\s*', '', kw_raw, flags=re.IGNORECASE).strip()
                set_run_font(kp.add_run("Index Terms"), size_pt=9, bold=True)
                set_run_font(kp.add_run("\u2014"),       size_pt=9, bold=True)
                set_run_font(kp.add_run(terms),          size_pt=9, italic=True)
            else:
                set_run_font(kp.add_run(kw_raw), size_pt=9, italic=True)

        # ── FIX: Section break — header=1 col, body will be 2 cols ──
        insert_section_break_single_to_double(doc)

        # ════════════════════════════════════════════════════
        # BODY — two column
        # ════════════════════════════════════════════════════
        diagram_index = 0   # track which uploaded diagram to embed next

        for raw in body_lines:
            line = clean_line(raw)
            if not line:
                continue

            # IEEE main heading
            if is_ieee_main_heading(line) or is_ieee_main_heading(line.upper()):
                normalised = re.sub(
                    r'^([ivx]+)\.\s+(.+)',
                    lambda m: m.group(1).upper() + '. ' + m.group(2).upper(),
                    line, flags=re.IGNORECASE
                )
                if not re.match(r'^[IVX]+\.', normalised):
                    normalised = line.upper()
                h = doc.add_paragraph()
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after  = Pt(3)
                set_run_font(h.add_run(normalised), size_pt=10, bold=True)

            # IEEE subsection A. Name
            elif paper_format == "IEEE" and is_ieee_subsection_heading(line):
                h = doc.add_paragraph()
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                h.paragraph_format.space_before = Pt(5)
                h.paragraph_format.space_after  = Pt(2)
                set_run_font(h.add_run(line), size_pt=10, bold=True, italic=True)

            # Generic heading for other formats
            elif is_generic_heading(line, body_headings):
                h = doc.add_paragraph()
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after  = Pt(3)
                set_run_font(h.add_run(line), size_pt=10, bold=True)

            # ── FIX: Figure placeholder — embed real diagram if available ──
            elif is_figure_placeholder(line):
                if diagram_index < len(diagram_images):
                    # Embed the actual uploaded image
                    fname, img_bytes, ext = diagram_images[diagram_index]
                    diagram_index += 1

                    # Figure paragraph with the actual image
                    fig_para = doc.add_paragraph()
                    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    fig_para.paragraph_format.space_before = Pt(8)
                    fig_para.paragraph_format.space_after  = Pt(2)

                    # Determine image type for docx
                    from docx.shared import Inches as DocxInches
                    img_stream = io.BytesIO(img_bytes)
                    try:
                        run_img = fig_para.add_run()
                        run_img.add_picture(img_stream, width=Inches(3.0))
                    except Exception as img_err:
                        print(f"Could not embed image: {img_err}")
                        # Fall back to placeholder text
                        set_run_font(fig_para.add_run(f"[ Figure {diagram_index}: {fname} ]"),
                                     size_pt=8, bold=True, italic=True)

                    # Caption below image
                    cap_para = doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_para.paragraph_format.space_before = Pt(2)
                    cap_para.paragraph_format.space_after  = Pt(8)
                    # Extract description from [DIAGRAM_HERE: description]
                    desc = re.sub(r'\[DIAGRAM_HERE[:\s]*', '', line, flags=re.IGNORECASE)
                    desc = desc.rstrip(']').strip()
                    caption_text = f"Fig. {diagram_index}. {desc}" if desc else f"Fig. {diagram_index}."
                    set_run_font(cap_para.add_run(caption_text), size_pt=8, italic=True)

                else:
                    # No uploaded image for this placeholder — show placeholder box
                    fp2 = doc.add_paragraph()
                    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    fp2.paragraph_format.space_before = Pt(6)
                    fp2.paragraph_format.space_after  = Pt(6)
                    desc = re.sub(r'\[DIAGRAM_HERE[:\s]*', '', line, flags=re.IGNORECASE).rstrip(']').strip()
                    placeholder_text = f"[ Figure: {desc} ]" if desc else "[ Figure: placeholder ]"
                    set_run_font(fp2.add_run(placeholder_text), size_pt=8, bold=True, italic=True)

            # Reference line [1] ...
            elif is_reference_line(line):
                rp = doc.add_paragraph()
                rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                rp.paragraph_format.left_indent       = Inches(0.3)
                rp.paragraph_format.first_line_indent = Inches(-0.3)
                rp.paragraph_format.space_before      = Pt(0)
                rp.paragraph_format.space_after       = Pt(2)
                set_run_font(rp.add_run(line), size_pt=8)

            # Equation line
            elif re.search(r'\(\d+\)\s*$', line) and len(line) < 120:
                ep = doc.add_paragraph()
                ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ep.paragraph_format.space_before = Pt(3)
                ep.paragraph_format.space_after  = Pt(3)
                set_run_font(ep.add_run(line), size_pt=10, italic=True)

            # Keywords repeated in body
            elif re.match(r'^(keywords?|index terms?)', line.lower()):
                kp2 = doc.add_paragraph()
                kp2.paragraph_format.space_after = Pt(4)
                set_run_font(kp2.add_run(line), size_pt=9, italic=True)

            # Normal body paragraph
            else:
                pp = doc.add_paragraph()
                pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pp.paragraph_format.first_line_indent = Inches(0.2)
                pp.paragraph_format.space_before      = Pt(0)
                pp.paragraph_format.space_after       = Pt(3)
                pp.paragraph_format.line_spacing      = Pt(12)
                set_run_font(pp.add_run(line), size_pt=10)

        # Apply two-column to the body section (last section)
        set_two_columns(doc.sections[-1])

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc_b64 = base64.b64encode(buf.read()).decode('utf-8')

        return jsonify({
            "success":     True,
            "content":     paper_content,
            "docx_base64": doc_b64,
            "filename":    f"{title[:50].replace(' ', '_')}_{paper_format}.docx"
        })

    except Exception as e:
        return jsonify({"error": f"Document generation failed: {str(e)}"}), 500


@app.route('/clarify-search', methods=['POST'])
def clarify_search():
    data  = request.get_json(silent=True) or {}
    query = data.get('query', '').strip()
    if not query:
        return jsonify({"error": "No query provided"}), 400

    prompt = f"""You are a research paper search assistant. A user has typed this search query:

"{query}"

Your job is to:
1. Extract the 4-6 most important academic keywords from this query
2. Generate exactly 3 short clarifying questions with 3-4 answer options each
   to help narrow down what research papers they actually need

The questions should help distinguish:
- The specific subtopic or angle they care about
- The type of paper (implementation, survey, theoretical, etc.)
- The domain or application area if ambiguous

Respond ONLY with valid JSON in this exact format, nothing else:
{{
  "keywords": ["keyword1", "keyword2", "keyword3", "keyword4"],
  "questions": [
    {{
      "question": "Question text here?",
      "options": ["Option A", "Option B", "Option C", "Option D"]
    }},
    {{
      "question": "Second question?",
      "options": ["Option A", "Option B", "Option C"]
    }},
    {{
      "question": "Third question?",
      "options": ["Option A", "Option B", "Option C"]
    }}
  ]
}}

Examples of good questions for "women safety app":
- "What aspect of safety are you focusing on?" → GPS tracking / SOS alerts / AI threat detection / Wearable devices
- "What type of paper do you need?" → Implementation/system design / Survey/review / Algorithm/model / User study
- "What platform or context?" → Mobile application / IoT/hardware / Social media analysis / General framework

Now generate for the query: "{query}"
Return ONLY the JSON, no explanation."""

    try:
        client   = Groq(api_key=os.environ.get('GROQ_API_KEY'))
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=600,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = response.choices[0].message.content.strip()

        # Strip markdown code fences if present
        raw = re.sub(r'^```(?:json)?\s*', '', raw)
        raw = re.sub(r'\s*```$',          '', raw)

        parsed = json.loads(raw)

        return jsonify({
            "success":   True,
            "keywords":  parsed.get("keywords",  []),
            "questions": parsed.get("questions", []),
        })

    except json.JSONDecodeError:
        # AI returned non-JSON — return empty so frontend falls back gracefully
        return jsonify({"success": True, "keywords": [], "questions": []})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(debug=False, port=8080)